"""
Document Converter - PDF, Word, and Excel Export
Converts Markdown documents to PDF, Word, and Excel formats.

Author: Akintade Akinokun
Purpose: CMMC Level 2 Gap Assessment Toolkit
"""

import os
import subprocess
import tempfile
from datetime import datetime
from typing import Optional, Tuple, List, Dict
import re


def check_dependencies() -> dict:
    """Check which conversion tools are available"""
    available = {
        "pandoc": False,
        "wkhtmltopdf": False,
        "python_docx": False,
        "weasyprint": False,
        "openpyxl": False
    }

    # Check for pandoc
    try:
        result = subprocess.run(["pandoc", "--version"], capture_output=True, timeout=5)
        available["pandoc"] = result.returncode == 0
    except:
        pass

    # Check for wkhtmltopdf
    try:
        result = subprocess.run(["wkhtmltopdf", "--version"], capture_output=True, timeout=5)
        available["wkhtmltopdf"] = result.returncode == 0
    except:
        pass

    # Check for python-docx
    try:
        import docx
        available["python_docx"] = True
    except ImportError:
        pass

    # Check for weasyprint
    try:
        import weasyprint
        available["weasyprint"] = True
    except ImportError:
        pass

    # Check for openpyxl (Excel)
    try:
        import openpyxl
        available["openpyxl"] = True
    except ImportError:
        pass

    return available


def markdown_to_html(markdown_content: str, title: str = "Document") -> str:
    """Convert markdown to styled HTML"""

    # Simple markdown to HTML conversion
    html = markdown_content

    # Convert headers
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', html, flags=re.MULTILINE)

    # Convert bold
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)

    # Convert italic
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)

    # Convert inline code
    html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)

    # Convert horizontal rules
    html = re.sub(r'^---+$', r'<hr>', html, flags=re.MULTILINE)

    # Convert tables
    lines = html.split('\n')
    in_table = False
    new_lines = []

    for i, line in enumerate(lines):
        if '|' in line and not line.strip().startswith('```'):
            if not in_table:
                new_lines.append('<table class="styled-table">')
                in_table = True

            # Skip separator lines
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue

            cells = [c.strip() for c in line.split('|')[1:-1]]

            # Check if this is a header row (first row of table)
            if in_table and len(new_lines) > 0 and new_lines[-1] == '<table class="styled-table">':
                row = '<tr>' + ''.join(f'<th>{c}</th>' for c in cells) + '</tr>'
            else:
                row = '<tr>' + ''.join(f'<td>{c}</td>' for c in cells) + '</tr>'
            new_lines.append(row)
        else:
            if in_table:
                new_lines.append('</table>')
                in_table = False

            # Convert list items
            if line.strip().startswith('- [ ]'):
                line = '<p class="checklist-item">☐ ' + line.strip()[5:].strip() + '</p>'
            elif line.strip().startswith('- [x]'):
                line = '<p class="checklist-item">☑ ' + line.strip()[5:].strip() + '</p>'
            elif line.strip().startswith('- '):
                line = '<li>' + line.strip()[2:] + '</li>'
            elif line.strip().startswith('* '):
                line = '<li>' + line.strip()[2:] + '</li>'
            elif re.match(r'^\d+\. ', line.strip()):
                line = '<li>' + re.sub(r'^\d+\. ', '', line.strip()) + '</li>'
            elif line.strip() and not line.strip().startswith('<'):
                line = '<p>' + line + '</p>'

            new_lines.append(line)

    if in_table:
        new_lines.append('</table>')

    html = '\n'.join(new_lines)

    # Wrap in full HTML document with styling
    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        @page {{
            size: letter;
            margin: 1in;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            font-size: 11pt;
            line-height: 1.5;
            color: #333;
            max-width: 8.5in;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{
            color: #1a5276;
            font-size: 24pt;
            border-bottom: 3px solid #1a5276;
            padding-bottom: 10px;
            margin-top: 30px;
        }}
        h2 {{
            color: #2c3e50;
            font-size: 16pt;
            border-bottom: 1px solid #3498db;
            padding-bottom: 5px;
            margin-top: 25px;
        }}
        h3 {{
            color: #2c3e50;
            font-size: 13pt;
            margin-top: 20px;
        }}
        h4 {{
            color: #555;
            font-size: 11pt;
            margin-top: 15px;
        }}
        .styled-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            font-size: 10pt;
        }}
        .styled-table th {{
            background-color: #2c3e50;
            color: white;
            padding: 10px 8px;
            text-align: left;
            font-weight: bold;
        }}
        .styled-table td {{
            padding: 8px;
            border-bottom: 1px solid #ddd;
        }}
        .styled-table tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}
        .styled-table tr:hover {{
            background-color: #e8f4f8;
        }}
        .checklist-item {{
            margin: 5px 0;
            padding: 5px 10px;
            background: #f8f9fa;
            border-left: 3px solid #3498db;
        }}
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            font-size: 10pt;
        }}
        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 20px 0;
        }}
        strong {{
            color: #2c3e50;
        }}
        li {{
            margin: 5px 0;
        }}
        p {{
            margin: 10px 0;
        }}
        .header-info {{
            background: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
    </style>
</head>
<body>
{html}
</body>
</html>
"""
    return full_html


def convert_to_pdf_with_pandoc(markdown_path: str, pdf_path: str) -> Tuple[bool, str]:
    """Convert markdown to PDF using pandoc"""
    try:
        cmd = [
            "pandoc",
            markdown_path,
            "-o", pdf_path,
            "--pdf-engine=wkhtmltopdf",
            "-V", "geometry:margin=1in",
            "-V", "fontsize=11pt"
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        if result.returncode == 0:
            return True, pdf_path
        else:
            # Try without wkhtmltopdf
            cmd = [
                "pandoc",
                markdown_path,
                "-o", pdf_path,
                "-V", "geometry:margin=1in"
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode == 0:
                return True, pdf_path
            return False, result.stderr
    except Exception as e:
        return False, str(e)


def convert_to_pdf_with_weasyprint(html_content: str, pdf_path: str) -> Tuple[bool, str]:
    """Convert HTML to PDF using WeasyPrint"""
    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(pdf_path)
        return True, pdf_path
    except Exception as e:
        return False, str(e)


def convert_to_pdf_builtin(markdown_path: str, pdf_path: str) -> Tuple[bool, str]:
    """
    Convert to PDF using built-in method (creates HTML and uses system print)
    Returns instructions if automatic conversion isn't possible
    """
    try:
        with open(markdown_path, 'r') as f:
            markdown_content = f.read()

        title = os.path.basename(markdown_path).replace('.md', '')
        html_content = markdown_to_html(markdown_content, title)

        # Save HTML file
        html_path = pdf_path.replace('.pdf', '.html')
        with open(html_path, 'w') as f:
            f.write(html_content)

        # Try weasyprint first
        deps = check_dependencies()

        if deps["weasyprint"]:
            success, result = convert_to_pdf_with_weasyprint(html_content, pdf_path)
            if success:
                return True, pdf_path

        if deps["pandoc"]:
            success, result = convert_to_pdf_with_pandoc(markdown_path, pdf_path)
            if success:
                return True, pdf_path

        # If no PDF converter available, return the HTML path
        return True, html_path

    except Exception as e:
        return False, str(e)


def convert_to_word(markdown_path: str, word_path: str) -> Tuple[bool, str]:
    """Convert markdown to Word document using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        with open(markdown_path, 'r') as f:
            content = f.read()

        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        lines = content.split('\n')
        in_table = False
        table_data = []

        for line in lines:
            line = line.rstrip()

            # Handle tables
            if '|' in line and not in_table:
                in_table = True
                table_data = []

            if in_table:
                if '|' in line:
                    if not re.match(r'^\|[\s\-:|]+\|$', line):
                        cells = [c.strip() for c in line.split('|')[1:-1]]
                        table_data.append(cells)
                else:
                    # End of table
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        for i, row_data in enumerate(table_data):
                            row = table.rows[i]
                            for j, cell_text in enumerate(row_data):
                                row.cells[j].text = cell_text
                                # Bold header row
                                if i == 0:
                                    row.cells[j].paragraphs[0].runs[0].bold = True
                        doc.add_paragraph()
                    in_table = False
                    table_data = []
                continue

            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            # Horizontal rule
            elif line.startswith('---'):
                doc.add_paragraph('_' * 50)
            # List items
            elif line.strip().startswith('- [ ]'):
                p = doc.add_paragraph('☐ ' + line.strip()[5:].strip(), style='List Bullet')
            elif line.strip().startswith('- [x]'):
                p = doc.add_paragraph('☑ ' + line.strip()[5:].strip(), style='List Bullet')
            elif line.strip().startswith('- '):
                p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif line.strip().startswith('* '):
                p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line.strip()):
                text = re.sub(r'^\d+\. ', '', line.strip())
                p = doc.add_paragraph(text, style='List Number')
            # Regular paragraph
            elif line.strip():
                p = doc.add_paragraph()
                # Handle bold text
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                # Empty line
                doc.add_paragraph()

        # Handle any remaining table
        if table_data:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            for i, row_data in enumerate(table_data):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    row.cells[j].text = cell_text

        doc.save(word_path)
        return True, word_path

    except ImportError:
        return False, "python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return False, str(e)


def convert_to_excel(company_name: str, output_path: str) -> Tuple[bool, str]:
    """
    Create an Excel workbook with Evidence Collection Checklist and Interview Guide.
    Features: dropdowns, conditional formatting, progress tracking.
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.worksheet.datavalidation import DataValidation
        from openpyxl.formatting.rule import FormulaRule, CellIsRule

        # Import evidence catalog and interview guide
        from evidence_checklist import EVIDENCE_CATALOG, INTERVIEW_GUIDE

        wb = Workbook()

        # ============================================
        # SHEET 1: Evidence Collection Checklist
        # ============================================
        ws_evidence = wb.active
        ws_evidence.title = "Evidence Checklist"

        # Styling
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        alt_row_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        wrap_alignment = Alignment(wrap_text=True, vertical='top')

        # Title row
        ws_evidence.merge_cells('A1:G1')
        ws_evidence['A1'] = f"CMMC Level 2 Evidence Collection Checklist - {company_name}"
        ws_evidence['A1'].font = Font(bold=True, size=16, color="1A5276")
        ws_evidence['A1'].alignment = Alignment(horizontal='center')
        ws_evidence.row_dimensions[1].height = 30

        # Info row
        ws_evidence.merge_cells('A2:G2')
        ws_evidence['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d')} | Total Artifacts: {sum(len(items) for items in EVIDENCE_CATALOG.values())}"
        ws_evidence['A2'].font = Font(italic=True, size=10)
        ws_evidence['A2'].alignment = Alignment(horizontal='center')

        # Progress row
        ws_evidence['A3'] = "Progress:"
        ws_evidence['B3'] = "=COUNTIF(E:E,\"Collected\")"
        ws_evidence['C3'] = "of"
        ws_evidence['D3'] = f"=COUNTA(E5:E500)-COUNTIF(E:E,\"N/A\")"
        ws_evidence['E3'] = "=IF(D3>0,ROUND(B3/D3*100,1)&\"%\",\"0%\")"
        ws_evidence['A3'].font = Font(bold=True)
        ws_evidence['E3'].font = Font(bold=True, color="2E7D32")

        # Headers
        headers = ["Category", "Artifact Name", "Description", "CMMC Controls", "Status", "Responsible Party", "Notes"]
        ws_evidence.append([])  # Empty row
        for col, header in enumerate(headers, 1):
            cell = ws_evidence.cell(row=5, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Column widths
        col_widths = [20, 35, 50, 20, 15, 20, 30]
        for i, width in enumerate(col_widths, 1):
            ws_evidence.column_dimensions[get_column_letter(i)].width = width

        # Freeze header row
        ws_evidence.freeze_panes = 'A6'

        # Status dropdown validation
        status_validation = DataValidation(
            type="list",
            formula1='"Pending,Collected,N/A"',
            allow_blank=True
        )
        status_validation.error = "Please select from the dropdown"
        status_validation.errorTitle = "Invalid Status"
        ws_evidence.add_data_validation(status_validation)

        # Add evidence items
        row = 6
        for category_id, category_data in EVIDENCE_CATALOG.items():
            category_name = category_data.get('category_name', category_id)
            for item in category_data.get('artifacts', []):
                ws_evidence.cell(row=row, column=1, value=category_name).border = border
                ws_evidence.cell(row=row, column=2, value=item.get('name', '')).border = border
                ws_evidence.cell(row=row, column=3, value=item.get('description', '')).border = border
                ws_evidence.cell(row=row, column=3).alignment = wrap_alignment
                controls = item.get('controls', [])
                ws_evidence.cell(row=row, column=4, value=', '.join(controls) if isinstance(controls, list) else str(controls)).border = border

                status_cell = ws_evidence.cell(row=row, column=5, value="Pending")
                status_cell.border = border
                status_validation.add(status_cell)

                ws_evidence.cell(row=row, column=6, value="").border = border
                ws_evidence.cell(row=row, column=7, value="").border = border

                # Alternate row coloring
                if row % 2 == 0:
                    for col in range(1, 8):
                        ws_evidence.cell(row=row, column=col).fill = alt_row_fill

                row += 1

        # Conditional formatting for status column
        green_fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")
        red_fill = PatternFill(start_color="FFCDD2", end_color="FFCDD2", fill_type="solid")
        gray_fill = PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")

        ws_evidence.conditional_formatting.add(
            f'E6:E{row}',
            CellIsRule(operator='equal', formula=['"Collected"'], fill=green_fill)
        )
        ws_evidence.conditional_formatting.add(
            f'E6:E{row}',
            CellIsRule(operator='equal', formula=['"Pending"'], fill=red_fill)
        )
        ws_evidence.conditional_formatting.add(
            f'E6:E{row}',
            CellIsRule(operator='equal', formula=['"N/A"'], fill=gray_fill)
        )

        # ============================================
        # SHEET 2: Interview Guide
        # ============================================
        ws_interview = wb.create_sheet("Interview Guide")

        # Title
        ws_interview.merge_cells('A1:E1')
        ws_interview['A1'] = f"CMMC Level 2 Interview Guide - {company_name}"
        ws_interview['A1'].font = Font(bold=True, size=16, color="1A5276")
        ws_interview['A1'].alignment = Alignment(horizontal='center')
        ws_interview.row_dimensions[1].height = 30

        # Calculate total questions
        total_questions = sum(
            sum(len(topic.get('questions', [])) for topic in role_data.get('topics', []))
            for role_data in INTERVIEW_GUIDE.values()
        )

        # Progress
        ws_interview['A2'] = "Completed:"
        ws_interview['B2'] = "=COUNTIF(E:E,\"Completed\")"
        ws_interview['C2'] = f"of {total_questions}"

        # Headers
        interview_headers = ["Role", "Topic", "Question", "Response Notes", "Status"]
        ws_interview.append([])
        ws_interview.append([])
        for col, header in enumerate(interview_headers, 1):
            cell = ws_interview.cell(row=5, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Column widths
        interview_widths = [25, 25, 55, 45, 15]
        for i, width in enumerate(interview_widths, 1):
            ws_interview.column_dimensions[get_column_letter(i)].width = width

        ws_interview.freeze_panes = 'A6'

        # Interview status dropdown
        interview_status = DataValidation(
            type="list",
            formula1='"Pending,Completed,Skipped"',
            allow_blank=True
        )
        ws_interview.add_data_validation(interview_status)

        # Add interview questions
        row = 6
        for role_id, role_data in INTERVIEW_GUIDE.items():
            role_name = role_data.get('role', role_id)
            for topic_data in role_data.get('topics', []):
                topic_name = topic_data.get('topic', '')
                for question in topic_data.get('questions', []):
                    ws_interview.cell(row=row, column=1, value=role_name).border = border
                    ws_interview.cell(row=row, column=2, value=topic_name).border = border
                    ws_interview.cell(row=row, column=3, value=question).border = border
                    ws_interview.cell(row=row, column=3).alignment = wrap_alignment
                    ws_interview.cell(row=row, column=4, value="").border = border

                    status_cell = ws_interview.cell(row=row, column=5, value="Pending")
                    status_cell.border = border
                    interview_status.add(status_cell)

                    if row % 2 == 0:
                        for col in range(1, 6):
                            ws_interview.cell(row=row, column=col).fill = alt_row_fill

                    row += 1

        # Conditional formatting for interview status (column E)
        ws_interview.conditional_formatting.add(
            f'E6:E{row}',
            CellIsRule(operator='equal', formula=['"Completed"'], fill=green_fill)
        )
        ws_interview.conditional_formatting.add(
            f'E6:E{row}',
            CellIsRule(operator='equal', formula=['"Pending"'], fill=red_fill)
        )

        # ============================================
        # SHEET 3: Control Assessment (110 Controls)
        # ============================================
        from cmmc_controls import CMMC_LEVEL2_CONTROLS

        ws_controls = wb.create_sheet("Control Assessment")

        # Title
        ws_controls.merge_cells('A1:H1')
        ws_controls['A1'] = f"CMMC Level 2 Control Assessment - {company_name}"
        ws_controls['A1'].font = Font(bold=True, size=16, color="1A5276")
        ws_controls['A1'].alignment = Alignment(horizontal='center')
        ws_controls.row_dimensions[1].height = 30

        # Summary row
        ws_controls.merge_cells('A2:H2')
        ws_controls['A2'] = "110 Security Requirements from NIST SP 800-171 Rev 2"
        ws_controls['A2'].font = Font(italic=True, size=10)
        ws_controls['A2'].alignment = Alignment(horizontal='center')

        # Progress tracking
        ws_controls['A3'] = "MET:"
        ws_controls['B3'] = '=COUNTIF(F:F,"MET")'
        ws_controls['C3'] = "NOT MET:"
        ws_controls['D3'] = '=COUNTIF(F:F,"NOT MET")'
        ws_controls['E3'] = "SPRS Score:"
        ws_controls['F3'] = '=110-COUNTIF(F:F,"NOT MET")*AVERAGE(G:G)'
        ws_controls['A3'].font = Font(bold=True)
        ws_controls['C3'].font = Font(bold=True)
        ws_controls['E3'].font = Font(bold=True)
        ws_controls['F3'].font = Font(bold=True, size=14, color="1A5276")

        # Headers
        control_headers = ["Domain", "Control ID", "Control Title", "Requirement", "Assessment Objective", "Status", "Point Value", "Notes"]
        for col, header in enumerate(control_headers, 1):
            cell = ws_controls.cell(row=5, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Column widths
        control_widths = [12, 15, 30, 50, 50, 12, 10, 30]
        for i, width in enumerate(control_widths, 1):
            ws_controls.column_dimensions[get_column_letter(i)].width = width

        ws_controls.freeze_panes = 'A6'

        # Control status dropdown
        control_status = DataValidation(
            type="list",
            formula1='"NOT ASSESSED,MET,NOT MET,N/A"',
            allow_blank=True
        )
        control_status.error = "Please select: MET, NOT MET, N/A, or NOT ASSESSED"
        ws_controls.add_data_validation(control_status)

        # Add all 110 controls
        row = 6
        for domain_id, domain_data in CMMC_LEVEL2_CONTROLS.items():
            domain_name = f"{domain_id}: {domain_data['domain_name']}"
            controls_dict = domain_data['controls']
            for control_id, control in controls_dict.items():
                ws_controls.cell(row=row, column=1, value=domain_name).border = border
                ws_controls.cell(row=row, column=2, value=control_id).border = border
                ws_controls.cell(row=row, column=2).font = Font(bold=True)
                ws_controls.cell(row=row, column=3, value=control.get('title', '')).border = border
                ws_controls.cell(row=row, column=3).alignment = wrap_alignment

                # Truncate requirement for readability
                requirement = control.get('requirement', '')[:200]
                if len(control.get('requirement', '')) > 200:
                    requirement += '...'
                ws_controls.cell(row=row, column=4, value=requirement).border = border
                ws_controls.cell(row=row, column=4).alignment = wrap_alignment

                # First assessment objective
                objectives = control.get('assessment_objectives', [])
                first_obj = objectives[0] if objectives else ''
                ws_controls.cell(row=row, column=5, value=first_obj).border = border
                ws_controls.cell(row=row, column=5).alignment = wrap_alignment

                status_cell = ws_controls.cell(row=row, column=6, value="NOT ASSESSED")
                status_cell.border = border
                control_status.add(status_cell)

                # Point value (for SPRS calculation) - default 1, some controls worth more
                ws_controls.cell(row=row, column=7, value=1).border = border
                ws_controls.cell(row=row, column=8, value="").border = border

                # Alternate row coloring
                if row % 2 == 0:
                    for col in range(1, 9):
                        ws_controls.cell(row=row, column=col).fill = alt_row_fill

                row += 1

        # Conditional formatting for control status
        ws_controls.conditional_formatting.add(
            f'F6:F{row}',
            CellIsRule(operator='equal', formula=['"MET"'], fill=green_fill)
        )
        ws_controls.conditional_formatting.add(
            f'F6:F{row}',
            CellIsRule(operator='equal', formula=['"NOT MET"'], fill=red_fill)
        )
        ws_controls.conditional_formatting.add(
            f'F6:F{row}',
            CellIsRule(operator='equal', formula=['"NOT ASSESSED"'], fill=PatternFill(start_color="FFF3E0", end_color="FFF3E0", fill_type="solid"))
        )

        # ============================================
        # SHEET 4: Summary Dashboard
        # ============================================
        ws_summary = wb.create_sheet("Dashboard")

        ws_summary['A1'] = f"CMMC Assessment Dashboard - {company_name}"
        ws_summary['A1'].font = Font(bold=True, size=18, color="1A5276")
        ws_summary.merge_cells('A1:D1')

        ws_summary['A3'] = "EVIDENCE COLLECTION"
        ws_summary['A3'].font = Font(bold=True, size=14)

        ws_summary['A4'] = "Total Artifacts:"
        ws_summary['B4'] = f"='Evidence Checklist'!D3"
        ws_summary['A5'] = "Collected:"
        ws_summary['B5'] = f"='Evidence Checklist'!B3"
        ws_summary['A6'] = "Progress:"
        ws_summary['B6'] = f"='Evidence Checklist'!E3"
        ws_summary['B6'].font = Font(bold=True, size=14, color="2E7D32")

        ws_summary['A8'] = "INTERVIEWS"
        ws_summary['A8'].font = Font(bold=True, size=14)

        ws_summary['A9'] = "Total Questions:"
        ws_summary['B9'] = total_questions
        ws_summary['A10'] = "Completed:"
        ws_summary['B10'] = "='Interview Guide'!B2"

        ws_summary['A12'] = "CONTROL ASSESSMENT"
        ws_summary['A12'].font = Font(bold=True, size=14)

        ws_summary['A13'] = "Total Controls:"
        ws_summary['B13'] = 110
        ws_summary['A14'] = "MET:"
        ws_summary['B14'] = "='Control Assessment'!B3"
        ws_summary['B14'].font = Font(bold=True, color="2E7D32")
        ws_summary['A15'] = "NOT MET:"
        ws_summary['B15'] = "='Control Assessment'!D3"
        ws_summary['B15'].font = Font(bold=True, color="C62828")
        ws_summary['A16'] = "Compliance %:"
        ws_summary['B16'] = '=IF(B14+B15>0,ROUND(B14/(B14+B15)*100,1)&"%","0%")'
        ws_summary['B16'].font = Font(bold=True, size=14, color="1A5276")

        ws_summary['A18'] = "ASSESSMENT INFO"
        ws_summary['A18'].font = Font(bold=True, size=14)
        ws_summary['A19'] = "Company:"
        ws_summary['B19'] = company_name
        ws_summary['A20'] = "Assessment Date:"
        ws_summary['B20'] = datetime.now().strftime('%Y-%m-%d')
        ws_summary['A21'] = "Assessor:"
        ws_summary['B21'] = ""

        ws_summary.column_dimensions['A'].width = 20
        ws_summary.column_dimensions['B'].width = 25

        # Save workbook
        wb.save(output_path)
        return True, output_path

    except ImportError:
        return False, "openpyxl not installed. Run: pip install openpyxl"
    except Exception as e:
        return False, str(e)


def generate_excel_checklist(company_name: str, output_dir: str = ".") -> Tuple[bool, str]:
    """
    Generate Excel checklist for a company assessment.

    Args:
        company_name: Name of the company being assessed
        output_dir: Directory to save the Excel file

    Returns:
        Tuple of (success, path_or_error)
    """
    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(' ', '_')
    output_path = os.path.join(output_dir, f"CMMC_Assessment_Checklist_{safe_name}.xlsx")
    return convert_to_excel(company_name, output_path)


def generate_controls_html(company_name: str, output_path: str) -> Tuple[bool, str]:
    """
    Generate an HTML visualization of all 110 CMMC controls for executive viewing.

    Args:
        company_name: Name of the company being assessed
        output_path: Path to save the HTML file

    Returns:
        Tuple of (success, path_or_error)
    """
    try:
        from cmmc_controls import CMMC_LEVEL2_CONTROLS, get_domain_summary

        domains = get_domain_summary()
        total_controls = sum(d['control_count'] for d in domains)

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CMMC Level 2 Control Assessment - {company_name}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            min-height: 100vh;
            color: #fff;
            padding: 20px;
        }}
        .container {{
            max-width: 1400px;
            margin: 0 auto;
        }}
        .header {{
            text-align: center;
            padding: 30px 0;
            border-bottom: 2px solid #3498db;
            margin-bottom: 30px;
        }}
        .header h1 {{
            font-size: 2.5em;
            color: #3498db;
            margin-bottom: 10px;
        }}
        .header .subtitle {{
            color: #888;
            font-size: 1.1em;
        }}
        .header .company {{
            color: #f39c12;
            font-size: 1.3em;
            margin-top: 10px;
        }}
        .summary-cards {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }}
        .card {{
            background: rgba(255,255,255,0.05);
            border-radius: 15px;
            padding: 25px;
            text-align: center;
            border: 1px solid rgba(255,255,255,0.1);
            transition: transform 0.3s, box-shadow 0.3s;
        }}
        .card:hover {{
            transform: translateY(-5px);
            box-shadow: 0 10px 30px rgba(0,0,0,0.3);
        }}
        .card .number {{
            font-size: 3em;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .card .label {{
            color: #888;
            text-transform: uppercase;
            font-size: 0.9em;
            letter-spacing: 1px;
        }}
        .card.total .number {{ color: #3498db; }}
        .card.domains .number {{ color: #9b59b6; }}
        .card.met .number {{ color: #2ecc71; }}
        .card.notmet .number {{ color: #e74c3c; }}

        .domains-section {{
            margin-top: 40px;
        }}
        .domains-section h2 {{
            color: #3498db;
            margin-bottom: 20px;
            font-size: 1.8em;
            border-bottom: 1px solid #3498db;
            padding-bottom: 10px;
        }}
        .domain-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
        }}
        .domain-card {{
            background: rgba(255,255,255,0.03);
            border-radius: 10px;
            overflow: hidden;
            border: 1px solid rgba(255,255,255,0.1);
        }}
        .domain-header {{
            background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%);
            padding: 15px 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .domain-header h3 {{
            font-size: 1em;
            color: #fff;
        }}
        .domain-header .count {{
            background: #3498db;
            padding: 5px 12px;
            border-radius: 20px;
            font-size: 0.9em;
            font-weight: bold;
        }}
        .domain-controls {{
            padding: 15px 20px;
            max-height: 300px;
            overflow-y: auto;
        }}
        .control-item {{
            padding: 10px 0;
            border-bottom: 1px solid rgba(255,255,255,0.05);
            display: flex;
            align-items: flex-start;
            gap: 10px;
        }}
        .control-item:last-child {{
            border-bottom: none;
        }}
        .control-id {{
            background: #2c3e50;
            padding: 3px 8px;
            border-radius: 4px;
            font-size: 0.75em;
            font-family: monospace;
            white-space: nowrap;
        }}
        .control-title {{
            color: #ccc;
            font-size: 0.9em;
            line-height: 1.4;
        }}
        .status-indicator {{
            width: 10px;
            height: 10px;
            border-radius: 50%;
            background: #f39c12;
            flex-shrink: 0;
            margin-top: 5px;
        }}
        .status-indicator.met {{ background: #2ecc71; }}
        .status-indicator.notmet {{ background: #e74c3c; }}

        .footer {{
            text-align: center;
            padding: 30px 0;
            margin-top: 40px;
            border-top: 1px solid rgba(255,255,255,0.1);
            color: #666;
        }}

        @media print {{
            body {{
                background: white;
                color: black;
            }}
            .card {{
                border: 1px solid #ddd;
            }}
            .card .number {{
                color: #333 !important;
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>CMMC Level 2 Control Assessment</h1>
            <div class="subtitle">Cybersecurity Maturity Model Certification | NIST SP 800-171 Rev 2</div>
            <div class="company">{company_name}</div>
        </div>

        <div class="summary-cards">
            <div class="card total">
                <div class="number">{total_controls}</div>
                <div class="label">Total Controls</div>
            </div>
            <div class="card domains">
                <div class="number">14</div>
                <div class="label">Security Domains</div>
            </div>
            <div class="card met">
                <div class="number" id="met-count">--</div>
                <div class="label">Controls MET</div>
            </div>
            <div class="card notmet">
                <div class="number" id="notmet-count">--</div>
                <div class="label">Gaps Identified</div>
            </div>
        </div>

        <div class="domains-section">
            <h2>Security Domains & Controls</h2>
            <div class="domain-grid">
"""

        # Add each domain
        for domain_id, domain_data in CMMC_LEVEL2_CONTROLS.items():
            controls_dict = domain_data['controls']
            control_count = len(controls_dict)
            html += f"""
                <div class="domain-card">
                    <div class="domain-header">
                        <h3>{domain_id}: {domain_data['domain_name']}</h3>
                        <span class="count">{control_count} controls</span>
                    </div>
                    <div class="domain-controls">
"""
            for control_id, control in controls_dict.items():
                html += f"""
                        <div class="control-item">
                            <span class="status-indicator"></span>
                            <span class="control-id">{control_id}</span>
                            <span class="control-title">{control.get('title', '')}</span>
                        </div>
"""
            html += """
                    </div>
                </div>
"""

        html += f"""
            </div>
        </div>

        <div class="footer">
            <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | CMMC Level 2 Gap Assessment Toolkit</p>
            <p style="margin-top: 10px; font-size: 0.9em;">
                This assessment covers all 110 security requirements from NIST SP 800-171 Rev 2<br>
                required for CMMC Level 2 certification.
            </p>
        </div>
    </div>
</body>
</html>
"""

        with open(output_path, 'w') as f:
            f.write(html)

        return True, output_path

    except Exception as e:
        return False, str(e)


def generate_controls_visualization(company_name: str, output_dir: str = ".") -> Dict[str, Tuple[bool, str]]:
    """
    Generate both Excel and HTML visualizations of 110 CMMC controls.

    Args:
        company_name: Name of the company being assessed
        output_dir: Directory to save files

    Returns:
        Dictionary with results for each format
    """
    results = {}
    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(' ', '_')

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Generate Excel
    excel_path = os.path.join(output_dir, f"CMMC_Assessment_Checklist_{safe_name}.xlsx")
    results['excel'] = convert_to_excel(company_name, excel_path)

    # Generate HTML
    html_path = os.path.join(output_dir, f"CMMC_Controls_Dashboard_{safe_name}.html")
    results['html'] = generate_controls_html(company_name, html_path)

    return results


def generate_documents_from_assessment(assessment, output_dir: str = ".") -> Dict[str, Tuple[bool, str]]:
    """
    Generate all documents from current assessment data with status filled in.

    This creates documents that reflect the current state of:
    - Evidence collection progress
    - Interview completion status
    - Control assessment status (MET/NOT MET)

    Args:
        assessment: CMMCAssessment object with current progress
        output_dir: Directory to save files

    Returns:
        Dictionary with results for each format generated
    """
    results = {}
    company_name = assessment.organization_name
    safe_name = re.sub(r'[^\w\s-]', '', company_name).replace(' ', '_')

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Generate Excel with current assessment data
    excel_path = os.path.join(output_dir, f"CMMC_Assessment_Checklist_{safe_name}.xlsx")
    results['excel'] = generate_excel_from_assessment(assessment, excel_path)

    # Generate HTML dashboard with current status
    html_path = os.path.join(output_dir, f"CMMC_Controls_Dashboard_{safe_name}.html")
    results['html'] = generate_html_from_assessment(assessment, html_path)

    # Generate Markdown/Text report
    md_path = os.path.join(output_dir, f"CMMC_Assessment_Report_{safe_name}.md")
    results['markdown'] = generate_markdown_from_assessment(assessment, md_path)

    # Generate Word document
    docx_path = os.path.join(output_dir, f"CMMC_Assessment_Report_{safe_name}.docx")
    results['word'] = generate_word_from_assessment(assessment, docx_path)

    return results


def generate_excel_from_assessment(assessment, output_path: str) -> Tuple[bool, str]:
    """Generate Excel workbook with current assessment status filled in."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.worksheet.datavalidation import DataValidation
        from openpyxl.formatting.rule import CellIsRule
        from cmmc_controls import CMMC_LEVEL2_CONTROLS

        company_name = assessment.organization_name
        wb = Workbook()

        # Styling
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        alt_row_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
        green_fill = PatternFill(start_color="C8E6C9", end_color="C8E6C9", fill_type="solid")
        red_fill = PatternFill(start_color="FFCDD2", end_color="FFCDD2", fill_type="solid")
        yellow_fill = PatternFill(start_color="FFF3E0", end_color="FFF3E0", fill_type="solid")
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        wrap_alignment = Alignment(wrap_text=True, vertical='top')

        # ============================================
        # SHEET 1: Evidence Collection with Status
        # ============================================
        ws_evidence = wb.active
        ws_evidence.title = "Evidence Checklist"

        ws_evidence.merge_cells('A1:G1')
        ws_evidence['A1'] = f"Evidence Collection Checklist - {company_name}"
        ws_evidence['A1'].font = Font(bold=True, size=16, color="1A5276")
        ws_evidence['A1'].alignment = Alignment(horizontal='center')

        ev_summary = assessment.get_evidence_summary()
        ws_evidence['A2'] = f"Progress: {ev_summary['collected']}/{ev_summary['total']} collected ({ev_summary['progress_percent']}%)"
        ws_evidence['A2'].font = Font(bold=True, color="2E7D32")

        headers = ["ID", "Category", "Artifact", "Description", "Status", "Responsible", "Notes"]
        for col, header in enumerate(headers, 1):
            cell = ws_evidence.cell(row=4, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border

        col_widths = [10, 20, 30, 40, 12, 20, 30]
        for i, width in enumerate(col_widths, 1):
            ws_evidence.column_dimensions[get_column_letter(i)].width = width

        ws_evidence.freeze_panes = 'A5'

        row = 5
        for artifact_id, data in assessment.evidence_status.items():
            ws_evidence.cell(row=row, column=1, value=artifact_id).border = border
            ws_evidence.cell(row=row, column=2, value=data.get('category', '')).border = border
            ws_evidence.cell(row=row, column=3, value=data.get('name', '')).border = border
            ws_evidence.cell(row=row, column=4, value=data.get('description', '')).border = border
            ws_evidence.cell(row=row, column=4).alignment = wrap_alignment

            status_cell = ws_evidence.cell(row=row, column=5, value=data.get('status', 'Pending'))
            status_cell.border = border
            if data.get('status') == 'Collected':
                status_cell.fill = green_fill
            elif data.get('status') == 'Pending':
                status_cell.fill = red_fill

            ws_evidence.cell(row=row, column=6, value=data.get('responsible_party', '')).border = border
            ws_evidence.cell(row=row, column=7, value=data.get('notes', '')).border = border

            if row % 2 == 0:
                for col in range(1, 8):
                    if ws_evidence.cell(row=row, column=col).fill == PatternFill():
                        ws_evidence.cell(row=row, column=col).fill = alt_row_fill
            row += 1

        # ============================================
        # SHEET 2: Interview Guide with Status
        # ============================================
        ws_interview = wb.create_sheet("Interview Guide")

        ws_interview.merge_cells('A1:F1')
        ws_interview['A1'] = f"Interview Guide - {company_name}"
        ws_interview['A1'].font = Font(bold=True, size=16, color="1A5276")

        int_summary = assessment.get_interview_summary()
        ws_interview['A2'] = f"Progress: {int_summary['completed']}/{int_summary['total']} completed ({int_summary['progress_percent']}%)"
        ws_interview['A2'].font = Font(bold=True, color="2E7D32")

        int_headers = ["ID", "Role", "Topic", "Question", "Status", "Response Notes"]
        for col, header in enumerate(int_headers, 1):
            cell = ws_interview.cell(row=4, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border

        int_widths = [10, 25, 20, 50, 12, 40]
        for i, width in enumerate(int_widths, 1):
            ws_interview.column_dimensions[get_column_letter(i)].width = width

        ws_interview.freeze_panes = 'A5'

        row = 5
        for q_id, data in assessment.interview_status.items():
            ws_interview.cell(row=row, column=1, value=q_id).border = border
            ws_interview.cell(row=row, column=2, value=data.get('role', '')).border = border
            ws_interview.cell(row=row, column=3, value=data.get('topic', '')).border = border
            ws_interview.cell(row=row, column=4, value=data.get('question', '')).border = border
            ws_interview.cell(row=row, column=4).alignment = wrap_alignment

            status_cell = ws_interview.cell(row=row, column=5, value=data.get('status', 'Pending'))
            status_cell.border = border
            if data.get('status') == 'Completed':
                status_cell.fill = green_fill
            elif data.get('status') == 'Pending':
                status_cell.fill = red_fill

            ws_interview.cell(row=row, column=6, value=data.get('response_notes', '')).border = border
            row += 1

        # ============================================
        # SHEET 3: Control Assessment with Status
        # ============================================
        ws_controls = wb.create_sheet("Control Assessment")

        ws_controls.merge_cells('A1:G1')
        ws_controls['A1'] = f"Control Assessment - {company_name}"
        ws_controls['A1'].font = Font(bold=True, size=16, color="1A5276")

        summary = assessment.get_assessment_summary()
        ws_controls['A2'] = f"Compliance: {summary['by_status']['MET']}/110 MET ({summary['compliance_score']}%)"
        ws_controls['A2'].font = Font(bold=True, color="2E7D32")

        ctrl_headers = ["Domain", "Control ID", "Title", "Status", "Implementation", "Gaps", "Notes"]
        for col, header in enumerate(ctrl_headers, 1):
            cell = ws_controls.cell(row=4, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border

        ctrl_widths = [15, 15, 35, 12, 40, 30, 30]
        for i, width in enumerate(ctrl_widths, 1):
            ws_controls.column_dimensions[get_column_letter(i)].width = width

        ws_controls.freeze_panes = 'A5'

        row = 5
        for domain_id, domain_data in CMMC_LEVEL2_CONTROLS.items():
            domain_name = f"{domain_id}: {domain_data['domain_name']}"
            for control_id, control in domain_data['controls'].items():
                ctrl_assessment = assessment.assessments.get(control_id)

                ws_controls.cell(row=row, column=1, value=domain_name).border = border
                ws_controls.cell(row=row, column=2, value=control_id).border = border
                ws_controls.cell(row=row, column=2).font = Font(bold=True)
                ws_controls.cell(row=row, column=3, value=control.get('title', '')).border = border
                ws_controls.cell(row=row, column=3).alignment = wrap_alignment

                status = ctrl_assessment.status if ctrl_assessment else 'NOT_ASSESSED'
                status_cell = ws_controls.cell(row=row, column=4, value=status)
                status_cell.border = border
                if status == 'MET':
                    status_cell.fill = green_fill
                elif status == 'NOT_MET':
                    status_cell.fill = red_fill
                elif status == 'NOT_ASSESSED':
                    status_cell.fill = yellow_fill

                impl = ctrl_assessment.implementation_description if ctrl_assessment else ''
                ws_controls.cell(row=row, column=5, value=impl).border = border
                ws_controls.cell(row=row, column=5).alignment = wrap_alignment

                gaps = ', '.join(ctrl_assessment.gaps_identified) if ctrl_assessment else ''
                ws_controls.cell(row=row, column=6, value=gaps).border = border

                notes = ctrl_assessment.assessor_notes if ctrl_assessment else ''
                ws_controls.cell(row=row, column=7, value=notes).border = border
                row += 1

        # ============================================
        # SHEET 4: Dashboard Summary
        # ============================================
        ws_dash = wb.create_sheet("Dashboard")

        ws_dash['A1'] = f"CMMC Assessment Dashboard - {company_name}"
        ws_dash['A1'].font = Font(bold=True, size=18, color="1A5276")
        ws_dash.merge_cells('A1:D1')

        ws_dash['A3'] = "OVERALL PROGRESS"
        ws_dash['A3'].font = Font(bold=True, size=14)

        ws_dash['A5'] = "Evidence Collection:"
        ws_dash['B5'] = f"{ev_summary['collected']}/{ev_summary['total']} ({ev_summary['progress_percent']}%)"
        ws_dash['A6'] = "Interviews:"
        ws_dash['B6'] = f"{int_summary['completed']}/{int_summary['total']} ({int_summary['progress_percent']}%)"
        ws_dash['A7'] = "Controls Assessed:"
        assessed = summary['by_status']['MET'] + summary['by_status']['NOT_MET']
        ws_dash['B7'] = f"{assessed}/110"

        ws_dash['A9'] = "COMPLIANCE STATUS"
        ws_dash['A9'].font = Font(bold=True, size=14)

        ws_dash['A10'] = "Controls MET:"
        ws_dash['B10'] = summary['by_status']['MET']
        ws_dash['B10'].font = Font(bold=True, color="2E7D32")
        ws_dash['A11'] = "Controls NOT MET:"
        ws_dash['B11'] = summary['by_status']['NOT_MET']
        ws_dash['B11'].font = Font(bold=True, color="C62828")
        ws_dash['A12'] = "Compliance Score:"
        ws_dash['B12'] = f"{summary['compliance_score']}%"
        ws_dash['B12'].font = Font(bold=True, size=14, color="1A5276")

        ws_dash['A14'] = "Generated:"
        ws_dash['B14'] = datetime.now().strftime('%Y-%m-%d %H:%M')

        ws_dash.column_dimensions['A'].width = 25
        ws_dash.column_dimensions['B'].width = 30

        wb.save(output_path)
        return True, output_path

    except ImportError:
        return False, "openpyxl not installed. Run: pip install openpyxl"
    except Exception as e:
        return False, str(e)


def generate_html_from_assessment(assessment, output_path: str) -> Tuple[bool, str]:
    """Generate HTML dashboard with current assessment status."""
    try:
        from cmmc_controls import CMMC_LEVEL2_CONTROLS

        company_name = assessment.organization_name
        summary = assessment.get_assessment_summary()
        ev_summary = assessment.get_evidence_summary()
        int_summary = assessment.get_interview_summary()

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CMMC Assessment Dashboard - {company_name}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
            min-height: 100vh;
            color: #fff;
            padding: 20px;
        }}
        .container {{ max-width: 1400px; margin: 0 auto; }}
        .header {{
            text-align: center;
            padding: 30px 0;
            border-bottom: 2px solid #3498db;
            margin-bottom: 30px;
        }}
        .header h1 {{ font-size: 2.5em; color: #3498db; margin-bottom: 10px; }}
        .header .company {{ color: #f39c12; font-size: 1.3em; margin-top: 10px; }}
        .summary-cards {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }}
        .card {{
            background: rgba(255,255,255,0.05);
            border-radius: 15px;
            padding: 25px;
            text-align: center;
            border: 1px solid rgba(255,255,255,0.1);
        }}
        .card .number {{ font-size: 2.5em; font-weight: bold; margin-bottom: 10px; }}
        .card .label {{ color: #888; text-transform: uppercase; font-size: 0.85em; }}
        .card.met .number {{ color: #2ecc71; }}
        .card.notmet .number {{ color: #e74c3c; }}
        .card.progress .number {{ color: #3498db; }}
        .card.evidence .number {{ color: #9b59b6; }}
        .card.interview .number {{ color: #f39c12; }}

        .domain-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin-top: 30px;
        }}
        .domain-card {{
            background: rgba(255,255,255,0.03);
            border-radius: 10px;
            overflow: hidden;
            border: 1px solid rgba(255,255,255,0.1);
        }}
        .domain-header {{
            background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%);
            padding: 15px 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .domain-header h3 {{ font-size: 1em; }}
        .domain-header .stats {{ display: flex; gap: 10px; }}
        .domain-header .met {{ color: #2ecc71; }}
        .domain-header .notmet {{ color: #e74c3c; }}
        .domain-controls {{ padding: 15px 20px; max-height: 250px; overflow-y: auto; }}
        .control-item {{
            padding: 8px 0;
            border-bottom: 1px solid rgba(255,255,255,0.05);
            display: flex;
            align-items: center;
            gap: 10px;
        }}
        .status-dot {{
            width: 10px;
            height: 10px;
            border-radius: 50%;
            flex-shrink: 0;
        }}
        .status-dot.met {{ background: #2ecc71; }}
        .status-dot.notmet {{ background: #e74c3c; }}
        .status-dot.pending {{ background: #f39c12; }}
        .control-id {{ font-family: monospace; font-size: 0.8em; color: #888; }}
        .control-title {{ color: #ccc; font-size: 0.9em; }}
        .footer {{
            text-align: center;
            padding: 30px 0;
            margin-top: 40px;
            border-top: 1px solid rgba(255,255,255,0.1);
            color: #666;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>CMMC Level 2 Assessment Dashboard</h1>
            <div class="company">{company_name}</div>
            <div style="color: #888; margin-top: 10px;">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
        </div>

        <div class="summary-cards">
            <div class="card met">
                <div class="number">{summary['by_status']['MET']}</div>
                <div class="label">Controls MET</div>
            </div>
            <div class="card notmet">
                <div class="number">{summary['by_status']['NOT_MET']}</div>
                <div class="label">Controls NOT MET</div>
            </div>
            <div class="card progress">
                <div class="number">{summary['compliance_score']}%</div>
                <div class="label">Compliance</div>
            </div>
            <div class="card evidence">
                <div class="number">{ev_summary['progress_percent']}%</div>
                <div class="label">Evidence Collected</div>
            </div>
            <div class="card interview">
                <div class="number">{int_summary['progress_percent']}%</div>
                <div class="label">Interviews Done</div>
            </div>
        </div>

        <h2 style="color: #3498db; border-bottom: 1px solid #3498db; padding-bottom: 10px;">Control Assessment by Domain</h2>
        <div class="domain-grid">
"""

        for domain_id, domain_data in CMMC_LEVEL2_CONTROLS.items():
            domain_met = 0
            domain_notmet = 0
            controls_html = ""

            for control_id, control in domain_data['controls'].items():
                ctrl_assessment = assessment.assessments.get(control_id)
                status = ctrl_assessment.status if ctrl_assessment else 'NOT_ASSESSED'

                if status == 'MET':
                    domain_met += 1
                    status_class = 'met'
                elif status == 'NOT_MET':
                    domain_notmet += 1
                    status_class = 'notmet'
                else:
                    status_class = 'pending'

                controls_html += f"""
                    <div class="control-item">
                        <span class="status-dot {status_class}"></span>
                        <span class="control-id">{control_id}</span>
                        <span class="control-title">{control.get('title', '')}</span>
                    </div>
"""

            html += f"""
            <div class="domain-card">
                <div class="domain-header">
                    <h3>{domain_id}: {domain_data['domain_name']}</h3>
                    <div class="stats">
                        <span class="met">✓ {domain_met}</span>
                        <span class="notmet">✗ {domain_notmet}</span>
                    </div>
                </div>
                <div class="domain-controls">
                    {controls_html}
                </div>
            </div>
"""

        html += f"""
        </div>

        <div class="footer">
            <p>CMMC Level 2 Gap Assessment Toolkit</p>
            <p style="margin-top: 10px; font-size: 0.9em;">
                110 Security Requirements | NIST SP 800-171 Rev 2
            </p>
        </div>
    </div>
</body>
</html>
"""

        with open(output_path, 'w') as f:
            f.write(html)

        return True, output_path

    except Exception as e:
        return False, str(e)


def generate_markdown_from_assessment(assessment, output_path: str) -> Tuple[bool, str]:
    """Generate Markdown report with current assessment status."""
    try:
        from cmmc_controls import CMMC_LEVEL2_CONTROLS

        company_name = assessment.organization_name
        summary = assessment.get_assessment_summary()
        ev_summary = assessment.get_evidence_summary()
        int_summary = assessment.get_interview_summary()

        md = f"""# CMMC Level 2 Assessment Report

**Organization:** {company_name}
**Assessment Date:** {assessment.assessment_date[:10]}
**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## Executive Summary

| Metric | Status |
|--------|--------|
| Controls MET | {summary['by_status']['MET']} of 110 |
| Controls NOT MET | {summary['by_status']['NOT_MET']} |
| Compliance Score | **{summary['compliance_score']}%** |
| Evidence Collected | {ev_summary['collected']} of {ev_summary['total']} ({ev_summary['progress_percent']}%) |
| Interviews Completed | {int_summary['completed']} of {int_summary['total']} ({int_summary['progress_percent']}%) |

---

## Control Assessment by Domain

"""

        for domain_id, domain_data in CMMC_LEVEL2_CONTROLS.items():
            domain_stats = summary['by_domain'].get(domain_id, {})
            md += f"""### {domain_id}: {domain_data['domain_name']}

**Status:** {domain_stats.get('met', 0)} MET | {domain_stats.get('not_met', 0)} NOT MET | {domain_stats.get('not_assessed', 0)} Not Assessed

| Control ID | Title | Status |
|------------|-------|--------|
"""
            for control_id, control in domain_data['controls'].items():
                ctrl_assessment = assessment.assessments.get(control_id)
                status = ctrl_assessment.status if ctrl_assessment else 'NOT_ASSESSED'
                status_icon = "✓" if status == "MET" else "✗" if status == "NOT_MET" else "○"
                md += f"| {control_id} | {control.get('title', '')} | {status_icon} {status} |\n"

            md += "\n"

        md += """---

## Evidence Collection Status

| Status | Count |
|--------|-------|
"""
        md += f"| Collected | {ev_summary['collected']} |\n"
        md += f"| Pending | {ev_summary['pending']} |\n"
        md += f"| N/A | {ev_summary['not_applicable']} |\n"

        md += """
---

## Interview Status

| Status | Count |
|--------|-------|
"""
        md += f"| Completed | {int_summary['completed']} |\n"
        md += f"| Pending | {int_summary['pending']} |\n"
        md += f"| Skipped | {int_summary['skipped']} |\n"

        md += f"""
---

*Report generated by CMMC Level 2 Gap Assessment Toolkit*
"""

        with open(output_path, 'w') as f:
            f.write(md)

        return True, output_path

    except Exception as e:
        return False, str(e)


def generate_word_from_assessment(assessment, output_path: str) -> Tuple[bool, str]:
    """Generate Word document from assessment data."""
    try:
        # First generate markdown
        md_path = output_path.replace('.docx', '_temp.md')
        success, _ = generate_markdown_from_assessment(assessment, md_path)

        if not success:
            return False, "Failed to generate markdown"

        # Convert markdown to Word
        success, result = convert_to_word(md_path, output_path)

        # Clean up temp file
        if os.path.exists(md_path):
            os.remove(md_path)

        return success, result

    except Exception as e:
        return False, str(e)


def convert_markdown_file(markdown_path: str, output_format: str = "both") -> dict:
    """
    Convert a markdown file to PDF and/or Word format.

    Args:
        markdown_path: Path to the markdown file
        output_format: "pdf", "word", or "both"

    Returns:
        Dictionary with results for each format
    """
    results = {}
    base_path = markdown_path.rsplit('.', 1)[0]

    if output_format in ["pdf", "both"]:
        pdf_path = base_path + ".pdf"
        success, result = convert_to_pdf_builtin(markdown_path, pdf_path)
        results["pdf"] = {
            "success": success,
            "path": result if success else None,
            "error": None if success else result
        }

    if output_format in ["word", "both"]:
        word_path = base_path + ".docx"
        success, result = convert_to_word(markdown_path, word_path)
        results["word"] = {
            "success": success,
            "path": result if success else None,
            "error": None if success else result
        }

    return results


def convert_all_documents(directory: str, output_format: str = "both") -> dict:
    """
    Convert all markdown files in a directory to PDF and Word.

    Args:
        directory: Directory containing markdown files
        output_format: "pdf", "word", or "both"

    Returns:
        Dictionary mapping filenames to conversion results
    """
    results = {}

    for filename in os.listdir(directory):
        if filename.endswith('.md'):
            filepath = os.path.join(directory, filename)
            results[filename] = convert_markdown_file(filepath, output_format)

    return results


def install_dependencies():
    """Provide instructions for installing conversion dependencies"""
    print("""
PDF, Word, and Excel Conversion Dependencies
============================================

For PDF generation, install ONE of these:

Option 1 - WeasyPrint (Recommended for Python):
    pip install weasyprint

Option 2 - Pandoc (Command-line tool):
    brew install pandoc        # macOS
    apt install pandoc         # Ubuntu/Debian
    choco install pandoc       # Windows

For Word generation:
    pip install python-docx

For Excel generation (Recommended):
    pip install openpyxl

Quick install all Python dependencies:
    pip install weasyprint python-docx openpyxl
""")


if __name__ == "__main__":
    print("Document Converter - PDF and Word Export")
    print("=" * 50)

    # Check dependencies
    deps = check_dependencies()
    print("\nAvailable conversion tools:")
    for tool, available in deps.items():
        status = "✓ Installed" if available else "✗ Not installed"
        print(f"  {tool}: {status}")

    if not any(deps.values()):
        print("\nNo conversion tools found.")
        install_dependencies()
    else:
        print("\nReady to convert documents.")

        # Demo conversion if evidence_request folder exists
        if os.path.exists("./evidence_request"):
            print("\nConverting documents in ./evidence_request/...")
            results = convert_all_documents("./evidence_request", "both")

            for filename, result in results.items():
                print(f"\n{filename}:")
                if "pdf" in result:
                    if result["pdf"]["success"]:
                        print(f"  PDF: ✓ {result['pdf']['path']}")
                    else:
                        print(f"  PDF: ✗ {result['pdf']['error']}")
                if "word" in result:
                    if result["word"]["success"]:
                        print(f"  Word: ✓ {result['word']['path']}")
                    else:
                        print(f"  Word: ✗ {result['word']['error']}")
